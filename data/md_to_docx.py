#!/usr/bin/env python3
"""Convert markdown to DOCX with basic styling."""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def md_to_docx(md_file, docx_file):
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines in tables
        if in_table and not line:
            i += 1
            continue
            
        # Detect table start
        if not in_table and '|' in line and i + 1 < len(lines) and '|' in lines[i + 1] and '-' in lines[i + 1]:
            in_table = True
            table_lines = [line]
            i += 1
            continue
        
        # Collect table lines
        if in_table:
            if '|' in line:
                table_lines.append(line)
                i += 1
                continue
            else:
                # End of table
                parse_table(doc, table_lines)
                in_table = False
                table_lines = []
                continue
        
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
        
        # Horizontal rules
        elif line.startswith('---'):
            doc.add_paragraph()
        
        # Bullet lists
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            p = doc.add_paragraph(text, style='List Bullet')
        
        # Numbered lists
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(text, style='List Number')
        
        # Checkboxes
        elif line.startswith('- [ ] ') or line.startswith('- [x] ') or line.startswith('- [X] '):
            checked = '[x]' in line.lower()
            text = re.sub(r'^- \[.\] ', '', line)
            p = doc.add_paragraph(('☑ ' if checked else '☐ ') + text)
        
        # Bold text
        elif '**' in line:
            p = doc.add_paragraph()
            parts = line.split('**')
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                if idx % 2 == 1:  # Odd indices are bold
                    run.bold = True
        
        # Regular paragraphs
        elif line:
            doc.add_paragraph(line)
        
        # Empty lines
        else:
            doc.add_paragraph()
        
        i += 1
    
    # Handle remaining table if file ends with one
    if in_table and table_lines:
        parse_table(doc, table_lines)
    
    doc.save(docx_file)
    print(f"✅ Converted {md_file} to {docx_file}")

def parse_table(doc, lines):
    """Parse markdown table and add to document."""
    if len(lines) < 2:
        return
    
    # Parse header
    headers = [cell.strip() for cell in lines[0].split('|') if cell.strip()]
    
    # Skip separator line
    # Parse data rows
    data_rows = []
    for line in lines[2:]:
        if '|' in line:
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                data_rows.append(cells)
    
    if not headers or not data_rows:
        return
    
    # Create table
    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    header_row = table.rows[0]
    for idx, header in enumerate(headers):
        if idx < len(header_row.cells):
            cell = header_row.cells[idx]
            cell.text = header
            # Bold headers
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
    
    # Add data
    for row_idx, data_row in enumerate(data_rows):
        table_row = table.rows[row_idx + 1]
        for col_idx, cell_data in enumerate(data_row):
            if col_idx < len(table_row.cells):
                table_row.cells[col_idx].text = cell_data
    
    doc.add_paragraph()  # Space after table

if __name__ == '__main__':
    md_to_docx('sc-operations-manual.md', 'sc-operations-manual.docx')
